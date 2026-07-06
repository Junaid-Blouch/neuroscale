import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.schemas.report_schema import NeuroScaleReportData

def generate_word_report(data: NeuroScaleReportData):
    # Naya blank document create karein
    doc = Document()
    
    # 1. Main Title
    title = doc.add_heading(data.report_title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 2. Executive Summary
    doc.add_heading("Executive Summary", level=1)
    
    p = doc.add_paragraph()
    p.add_run("Date Range: ").bold = True
    p.add_run(f"{data.start_date} to {data.end_date}\n")
    
    # 3. Key Metrics (With Colors)
    doc.add_heading("System Performance & FinOps Metrics", level=2)
    kpi_p = doc.add_paragraph()
    
    kpi_p.add_run("Total Cost Saved: ").bold = True
    cost_run = kpi_p.add_run(f"${data.total_cost_saved} USD\n")
    cost_run.font.color.rgb = RGBColor(0, 128, 0) # Green Color
    
    kpi_p.add_run("Overloads Prevented: ").bold = True
    kpi_p.add_run(f"{data.overload_prevented_events} Events\n")
    
    kpi_p.add_run("Model Drift Status: ").bold = True
    drift_run = kpi_p.add_run(f"{data.drift_status}\n")
    
    # Adaptive Color based on status
    if "Stable" in data.drift_status or "Normal" in data.drift_status:
        drift_run.font.color.rgb = RGBColor(0, 128, 0) # Green
    else:
        drift_run.font.color.rgb = RGBColor(255, 0, 0) # Red
        
    # 4. Data Table (LSTM & PPO Logs)
    doc.add_heading("LSTM Workload Forecast & PPO Scaling Logs", level=2)
    
    # Ek khoobsurat table banayen
    table = doc.add_table(rows=1, cols=3)
    table.style = "Medium Shading 1 Accent 1" # Word ka built-in professional design
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Timestamp"
    hdr_cells[1].text = "LSTM Predicted CPU (%)"
    hdr_cells[2].text = "PPO Target Replicas"
    
    # Table mein data insert karein
    for item in data.metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.timestamp)
        row_cells[1].text = f"{item.lstm_predicted_cpu}%"
        row_cells[2].text = str(item.ppo_target_replicas)
        
    # 5. File Save Karein
    output_path = "NeuroScale_Automated_Report.docx"
    doc.save(output_path)
    
    return output_path
