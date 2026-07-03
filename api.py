from fastapi import FastAPI, UploadFile, File, Query
from fastapi.responses import FileResponse, JSONResponse
from enum import Enum
import pandas as pd
import numpy as np
import torch
import torch.nn as nn
import joblib
import os
import datetime
from stable_baselines3 import PPO

# --- NEW: Graph Generation Libraries ---
import matplotlib.pyplot as plt
import io
import base64
# ---------------------------------------

# Report libraries
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from docx import Document
import xlsxwriter

app = FastAPI(title="NeuroScale Enterprise AI")

BASE_PATH = os.path.dirname(os.path.abspath(__file__))
MODEL_PATH = os.path.join(BASE_PATH, "models", "trained_lstm.pth")
SCALER_PATH = os.path.join(BASE_PATH, "models", "workload_scaler.save")
PPO_PATH = os.path.join(BASE_PATH, "models", "ppo_cloud_agent.zip")

MAX_SERVERS = 20
SERVER_CAPACITY = 1.0 / MAX_SERVERS

class ReportFormat(str, Enum):
    json = "json"
    pdf = "pdf"
    word = "word"
    excel = "excel"

# LSTM Model
class LSTMModel(nn.Module):
    def __init__(self):
        super().__init__()
        self.lstm = nn.LSTM(3, 64, 1, batch_first=True)
        self.fc = nn.Linear(64, 1)

    def forward(self, x):
        out, _ = self.lstm(x)
        return self.fc(out[:, -1, :])

lstm_model = LSTMModel()
lstm_model.load_state_dict(torch.load(MODEL_PATH))
lstm_model.eval()

scaler = joblib.load(SCALER_PATH)
ppo_model = PPO.load(PPO_PATH)

def multi_step_forecast(initial_scaled_window, steps=168):
    predictions = []
    current_seq = initial_scaled_window.copy()

    for _ in range(steps):
        tensor = torch.tensor(
            current_seq.reshape(1, len(current_seq), 3),
            dtype=torch.float32
        )
        with torch.no_grad():
            pred_scaled = lstm_model(tensor)

        temp = np.zeros((1,3))
        temp[0,0] = pred_scaled.item()
        pred = scaler.inverse_transform(temp)[0][0]

        predictions.append(pred)

        new_row = current_seq[-1].copy()
        new_row[0] = pred_scaled.item()
        current_seq = np.vstack([current_seq[1:], new_row])

    return predictions

# --- NEW: UI Server Endpoint ---
@app.get("/")
def serve_frontend():
    # Yeh frontend ko browser mein show karega
    return FileResponse("web/index.html")
# -------------------------------

@app.post("/predict")
async def predict(
    file: UploadFile = File(...),
    format: ReportFormat = Query(ReportFormat.json)
):

    df = pd.read_csv(file.file)

    if len(df) < 60:
        return {"error": "Minimum 60 rows required"}

    workload = df.iloc[:,0].values
    workload = (workload - workload.min()) / (workload.max() - workload.min())

    last = workload[-60:]

    df_temp = pd.DataFrame()
    df_temp["workload"] = last
    df_temp["workload"] = df_temp["workload"].rolling(5).mean()
    df_temp.dropna(inplace=True)
    df_temp["rolling_mean"] = df_temp["workload"].rolling(20).mean()
    df_temp["rolling_std"] = df_temp["workload"].rolling(20).std()
    df_temp.dropna(inplace=True)

    features = df_temp[["workload","rolling_mean","rolling_std"]].values
    scaled = scaler.transform(features)

    forecast = multi_step_forecast(scaled, 168)

    start_time = datetime.datetime.now()

    report_data = []
    active_servers = 5

    for i in range(168):

        predicted_load = forecast[i]
        load_delta = 0 if i == 0 else forecast[i] - forecast[i-1]

        state = np.array([predicted_load, predicted_load, load_delta, active_servers], dtype=np.float32)
        action, _ = ppo_model.predict(state)

        if action == 0:
            active_servers = max(1, active_servers - 1)
            decision = "Decrease"
        elif action == 2:
            active_servers = min(MAX_SERVERS, active_servers + 1)
            decision = "Increase"
        else:
            decision = "Keep"

        report_data.append([
            (start_time + datetime.timedelta(hours=i)).strftime("%Y-%m-%d %H:%M"),
            round(predicted_load,4),
            decision,
            active_servers
        ])

    if format == ReportFormat.json:
        # --- NEW: UI Integration Logic (Graph & Dictionary Formatting) ---
        
        # 1. Generate Dark-Mode Graph
        plt.figure(figsize=(10, 4))
        loads = [row[1] for row in report_data]
        plt.plot(range(len(loads)), loads, color='#00ffcc', linewidth=2)
        
        # Styling to match your SaaS UI
        plt.gca().set_facecolor('#050505')
        plt.gcf().patch.set_facecolor('#050505')
        plt.tick_params(colors='white')
        plt.grid(color='rgba(255, 255, 255, 0.1)', linestyle='--', alpha=0.5)
        plt.tight_layout()

        # Save to buffer and encode
        buf = io.BytesIO()
        plt.savefig(buf, format="png", bbox_inches='tight')
        buf.seek(0)
        graph_base64 = base64.b64encode(buf.read()).decode("utf-8")
        plt.close()

        # 2. Format Data for the Frontend Table
        formatted_forecast = []
        for row in report_data:
            formatted_forecast.append({
                "datetime": row[0],
                "predicted_load": row[1],
                "action": row[2],
                "servers": row[3]
            })

        return JSONResponse(content={
            "forecast": formatted_forecast,
            "forecast_graph_base64": graph_base64
        })
        # ----------------------------------------------------------------

    if format == ReportFormat.pdf:

        pdf_path = "NeuroScale_Report.pdf"
        doc = SimpleDocTemplate(pdf_path, pagesize=A4)

        pdfmetrics.registerFont(TTFont('TimesNewRoman', 'C:/Windows/Fonts/times.ttf'))

        heading = ParagraphStyle(name='Heading', fontName='TimesNewRoman', fontSize=16)
        elements = []

        elements.append(Paragraph("<b>NeuroScale Weekly Forecast Report</b>", heading))
        elements.append(Spacer(1, 12))

        table_data = [["Date-Time","Predicted Load","Action","Servers"]]
        table_data.extend(report_data[:50])

        table = Table(table_data)
        table.setStyle([('GRID',(0,0),(-1,-1),1,colors.black)])
        elements.append(table)

        doc.build(elements)

        return FileResponse(pdf_path, filename="NeuroScale_Report.pdf")

    if format == ReportFormat.word:

        doc = Document()
        doc.add_heading("NeuroScale Weekly Forecast Report", level=1)

        for row in report_data[:50]:
            doc.add_paragraph(f"{row[0]} | Load: {row[1]} | Action: {row[2]} | Servers: {row[3]}")

        word_path = "NeuroScale_Report.docx"
        doc.save(word_path)

        return FileResponse(word_path, filename="NeuroScale_Report.docx")

    if format == ReportFormat.excel:

        excel_path = "NeuroScale_Report.xlsx"
        workbook = xlsxwriter.Workbook(excel_path)
        worksheet = workbook.add_worksheet("Forecast")

        headers = ["Date-Time","Predicted Load","Action","Servers"]

        for col, header in enumerate(headers):
            worksheet.write(0,col,header)

        for row_num,row in enumerate(report_data[:50], start=1):
            for col_num,cell in enumerate(row):
                worksheet.write(row_num,col_num,cell)

        workbook.close()

        return FileResponse(excel_path, filename="NeuroScale_Report.xlsx")

    return {"error": "Invalid format"}